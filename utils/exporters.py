"""Export utilities: PDF, DOCX, JSON.
"""
from __future__ import annotations
from typing import List, Dict
from io import BytesIO

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
from docx import Document
import json

def build_pdf(summary: str, mcqs: List[Dict], flashcards: List[Dict]) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 40
    y = height - margin

    def draw_heading(text):
        nonlocal y
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, text)
        y -= 24

    def draw_paragraph(text, font_size=10):
        nonlocal y
        c.setFont("Helvetica", font_size)
        wrapped = simpleSplit(text, "Helvetica", font_size, width - 2*margin)
        for line in wrapped:
            if y < margin + 40:
                c.showPage()
                y = height - margin
                c.setFont("Helvetica", font_size)
            c.drawString(margin, y, line)
            y -= 14
        y -= 6

    draw_heading("Summarized Notes")
    draw_paragraph(summary or "(No summary)")

    draw_heading("Multiple Choice Questions")
    if not mcqs:
        draw_paragraph("(No MCQs)")
    else:
        for i, q in enumerate(mcqs, 1):
            draw_paragraph(f"{i}. {q['question']}")
            for opt in q['options']:
                draw_paragraph(f" - {opt}", font_size=9)
            draw_paragraph(f"Answer: {q['answer']}", font_size=9)

    draw_heading("Flashcards")
    if not flashcards:
        draw_paragraph("(No flashcards)")
    else:
        for fc in flashcards:
            draw_paragraph(f"Term: {fc['term']}")
            draw_paragraph(f"Definition: {fc['definition']}", font_size=9)

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.read()


def build_docx(summary: str, mcqs: List[Dict], flashcards: List[Dict]) -> bytes:
    doc = Document()
    doc.add_heading('Summarized Notes', level=1)
    doc.add_paragraph(summary or "(No summary)")

    doc.add_heading('Multiple Choice Questions', level=1)
    if not mcqs:
        doc.add_paragraph('(No MCQs)')
    else:
        for i, q in enumerate(mcqs, 1):
            doc.add_paragraph(f"{i}. {q['question']}")
            for opt in q['options']:
                doc.add_paragraph(f" - {opt}")
            doc.add_paragraph(f"Answer: {q['answer']}")

    doc.add_heading('Flashcards', level=1)
    if not flashcards:
        doc.add_paragraph('(No flashcards)')
    else:
        for fc in flashcards:
            doc.add_paragraph(f"Term: {fc['term']}")
            doc.add_paragraph(f"Definition: {fc['definition']}")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def build_flashcards_json(flashcards: List[Dict]) -> bytes:
    return json.dumps(flashcards, indent=2).encode('utf-8')

__all__ = ["build_pdf", "build_docx", "build_flashcards_json"]
